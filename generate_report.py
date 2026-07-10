import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    doc = docx.Document()
    
    # Set default font to resemble LaTeX's default serif (similar to Times New Roman)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # PAGE 1: Cover Page
    doc.add_paragraph('\n\n\n\n')
    title = doc.add_paragraph('Design and Development of a Modern MERN Stack E-Commerce Platform')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(18)
        run.bold = True
    
    doc.add_paragraph('\n\n')
    sub = doc.add_paragraph('An Internship Report\nSubmitted in partial fulfillment of the requirements for the Degree of\nBachelor of Science in Computer Science and Engineering')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(12)
    
    doc.add_paragraph('\n\n')
    sub = doc.add_paragraph('Submitted by\n\nRaisa Tabassum Kabir     [Your ID]')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        if 'Raisa' in run.text or '[Your ID]' in run.text:
            run.bold = True
            
    doc.add_paragraph('\n\n')
    sub = doc.add_paragraph('Supervised by\n\n[Supervisor Name]\nLecturer\nDepartment of Computer Science and Engineering\nSoutheast University')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        if 'Supervisor Name' in run.text:
            run.bold = True

    doc.add_paragraph('\n\n\n')
    sub = doc.add_paragraph('Department of Computer Science and Engineering\nSoutheast University\nDhaka, Bangladesh\n\nJuly 2026')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.bold = True

    doc.add_page_break()
    
    # PAGE 2: Letter of Transmittal
    h = doc.add_heading('Letter of Transmittal', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('July 06, 2026\n\nThe Chairman,\nDepartment of Computer Science and Engineering\nSoutheast University\nTejgaon, Dhaka, Bangladesh\n\nThrough: Supervisor, [Supervisor Name]\n\nSubject: Submission of Internship Report.\n\nDear Sir,\nWith due respect, I am pleased to submit my internship report entitled "Design and Development of a Modern MERN Stack E-Commerce Platform" in partial fulfillment of the requirements for completing the internship program.\n\nDuring my internship, I worked on developing an end-to-end e-commerce platform using the MERN stack (MongoDB, Express, React, Node.js). The system combines secure Stripe payment processing, real-time Socket.IO notifications, Cloudinary image storage, and a robust admin dashboard.\n\nThis report provides an overview of the tasks I performed, the methodologies I followed, the tools and technologies I used, and the key learnings gained throughout the internship. I have prepared this report carefully according to the given instructions and requirements.\n\nThank you for your support and consideration.\n\nSincerely Yours,\n\n\nRaisa Tabassum Kabir\n[Your ID]\n\nSupervisor:\n\n_______________________\n[Supervisor Name]\nLecturer & Supervisor\nDepartment of Computer Science and Engineering\nSoutheast University')
    doc.add_page_break()

    # PAGE 3: Candidate's Declaration
    h = doc.add_heading('CANDIDATE\'S DECLARATION', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('I, hereby, declare that the thesis presented in this report is the outcome of the investigation performed by us under the supervision of [Supervisor Name], Lecturer, Department of Computer Science and Engineering, Southeast University. The work was done through CSE489: Internship course, in accordance with the course curriculum of the Department for the Bachelor of Science in Computer Science and Engineering program.\n\nIt is also declared that neither this research nor any part thereof has been submitted anywhere else for the award of any degree, diploma or other qualifications.\n\n\n\n_______________________\nRaisa Tabassum Kabir\n[Your ID]')
    doc.add_page_break()

    # PAGE 4: Certification
    h = doc.add_heading('CERTIFICATION', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('This report titled, "Design and Development of a Modern MERN Stack E-Commerce Platform", submitted by the people mentioned below has been accepted as satisfactory in partial fulfillment of the requirements for the degree B.Sc. in Computer Science and Engineering in July 2026.\n\nMember:\n\nRaisa Tabassum Kabir     [Your ID]\n\n\n\n_______________________\n[Supervisor Name]\nLecturer & Supervisor\nDepartment of Computer Science and Engineering\nSoutheast University\n\n\n_______________________\nShahriar Manzoor\nAssociate Professor & Chairman\nDepartment of Computer Science and Engineering\nSoutheast University')
    doc.add_page_break()

    # PAGE 5: Executive Summary
    h = doc.add_heading('Executive Summary', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('This report summarizes my internship experience, where I contributed to the design and development of a modern, full-stack E-Commerce platform.\n\nThe system replaces traditional manual store management with an automated, end-to-end digital pipeline capable of handling products, orders, users, and secure payments. It integrates a React and Tailwind CSS frontend for a seamless user experience, with a Node.js and Express backend. MongoDB was utilized for scalable data storage, while Stripe was integrated for secure credit card processing.\n\nI also implemented real-time notifications via Socket.IO and robust image storage using Cloudinary and Multer. Overall, the internship strengthened my skills in full-stack web development, system design, API integration, and building secure, performance-aware pipelines for real-world applications.')
    doc.add_page_break()

    # PAGE 6: Acknowledgement
    h = doc.add_heading('ACKNOWLEDGEMENT', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('I would like to express my sincere gratitude to [Supervisor Name], Lecturer, for her valuable guidance, constructive feedback and continuous encouragement throughout my internship work titled Design and Development of a Modern MERN Stack E-Commerce Platform. Her insights and suggestions played a vital role in shaping the direction and quality of this work.\n\nIt is my privilege to extend my humble thanks to Southeast University for providing me with the facilities and support that contributed to the preparation of this internship report. Without the encouragement, inspiration and support of many people, this report would not have been successfully completed.\n\nI would like to acknowledge the contributions of everyone who supported me directly or indirectly during this internship, including faculty members, colleagues and peers, for their encouragement and helpful discussions.\n\nFinally, I am grateful to my family and friends for their unwavering support, patience and motivation throughout this journey.\n\n\nDhaka\nJuly 06, 2026\n\n\nRaisa Tabassum Kabir')
    doc.add_page_break()

    # PAGE 7: Table of Contents
    h = doc.add_heading('Contents', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph('LETTER OF TRANSMITTAL\nCANDIDATES\' DECLARATION\nCERTIFICATION\nEXECUTIVE SUMMARY\nACKNOWLEDGEMENT\n1 Introduction\n2 Background Study\n3 Details of Internship Work\n4 Key Learnings\n5 Limitations and Recommendations\n6 Conclusion\nReferences')
    doc.add_page_break()
    
    # Chapter 1
    doc.add_heading('Chapter 1\n\nIntroduction', level=1)
    doc.add_heading('1.1 Introduction', level=2)
    doc.add_paragraph('This internship report is prepared as a partial fulfillment of the requirements of my academic program. As part of the internship experience, I was engaged in a technology-oriented project to build a full-stack E-Commerce platform using the MERN stack.\n\nThe task was devoted to creating a robust digital marketplace that supports user authentication, product catalogs, shopping carts, secure payments, and a comprehensive admin dashboard.')
    doc.add_heading('1.2 Objective of the Internship', level=2)
    doc.add_paragraph('The objective of the internship project was to gain practical experience in modern web development and contribute to digital transformation. The purpose of this study is:\n- To develop a scalable e-commerce backend using Node.js and MongoDB.\n- To design a responsive frontend using React and Tailwind CSS.\n- To integrate third-party services like Stripe for payments and Cloudinary for media storage.')
    doc.add_page_break()
    
    # Chapter 2
    doc.add_heading('Chapter 2\n\nBackground Study', level=1)
    doc.add_heading('2.1 Literature Review', level=2)
    doc.add_paragraph('Before the widespread adoption of modern web frameworks, building an e-commerce platform required managing complex server-side rendered pages and maintaining challenging state transitions. Traditional monolithic architectures often struggled with scalability and real-time updates.\n\nThe MERN stack (MongoDB, Express.js, React, Node.js) represents a paradigm shift towards decoupled, API-driven development. Research into modern web architectures highlights the benefits of using a single language (JavaScript) across the entire stack, which accelerates development and improves maintainability. Furthermore, integrating specialized cloud services (e.g., Stripe for payments, Cloudinary for storage) reduces the burden on the core application and enhances security.')
    doc.add_page_break()

    # Chapter 3
    doc.add_heading('Chapter 3\n\nDetails of Internship Work', level=1)
    doc.add_heading('3.1 Working Methodology', level=2)
    doc.add_paragraph('The project was executed in a chronological and systematic manner. Initially, I conducted research on the general process of e-commerce workflows, including cart management, user authentication (JWT), and payment gateways. Then, I set up the foundation of the backend API using Express.js and MongoDB (Mongoose).\n\nFor the frontend, Vite and React were used to bootstrap the application. Zustand was chosen for state management to ensure seamless synchronization of the user\'s shopping cart and authentication state across components. Real-time features, such as order updates, were integrated using Socket.IO.')
    doc.add_page_break()

    # Chapter 4
    doc.add_heading('Chapter 4\n\nKey Learnings', level=1)
    doc.add_heading('4.1 Key Learning', level=2)
    doc.add_paragraph('1. Full-Stack Development:\n- Acquired specialist and theoretical knowledge on developing MERN stack applications.\n- Gained practical knowledge on building RESTful APIs and securing them with JWT.\n\n2. Third-Party Integrations:\n- Learned to securely process payments using Stripe Elements.\n- Implemented image uploads via Multer and Cloudinary.\n- Integrated Google\'s Gmail REST API for sending transactional emails directly over HTTPS.\n\n3. Real-Time Communication:\n- Acquired practical experience with Socket.IO to push real-time updates to connected clients.')
    doc.add_page_break()
    
    # Chapter 5
    doc.add_heading('Chapter 5\n\nLimitations and Recommendations', level=1)
    doc.add_heading('5.1 Limitations', level=2)
    doc.add_paragraph('Although the proposed framework is fully functional, it presents a number of limitations. The current system uses local state and Zustand for cart management, which might not sync across multiple devices if the user logs in from a different browser before checking out. Furthermore, the search functionality is based on basic database queries and might not scale efficiently with a massive product catalog.')
    doc.add_heading('5.2 Recommendations for Future Improvement', level=2)
    doc.add_paragraph('- Implement advanced search capabilities using Elasticsearch or Algolia.\n- Transition cart state to be fully database-backed for cross-device synchronization.\n- Add a recommendation engine based on user browsing and purchase history.')
    doc.add_page_break()

    # Chapter 6
    doc.add_heading('Chapter 6\n\nConclusion', level=1)
    doc.add_paragraph('The internship was a highly useful experience in the field of modern web application development. The core goal of the project was to design and execute an end-to-end e-commerce pipeline, which was successfully achieved by building a robust MERN stack application.\n\nIn addition to technical implementation, this internship made me better aware of operational issues related to digital commerce, including data security, compliance with payment standards, and reliable document storage. The experience has greatly enhanced my capacity to solve problems and my willingness to work on massive digital transformation projects.')
    doc.add_page_break()
    
    # References
    doc.add_heading('References', level=1)
    doc.add_paragraph('[1] "React - A JavaScript library for building user interfaces." https://reactjs.org/. Accessed 2026.\n[2] "Node.js." https://nodejs.org/en/. Accessed 2026.\n[3] "MongoDB: The Developer Data Platform." https://www.mongodb.com/. Accessed 2026.')
    
    # Set heading fonts to Times New Roman and black
    for style in doc.styles:
        if style.name.startswith('Heading'):
            font = style.font
            font.name = 'Times New Roman'
            font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    
    doc.save('Internship_Report_ECommerce.docx')

if __name__ == "__main__":
    main()
